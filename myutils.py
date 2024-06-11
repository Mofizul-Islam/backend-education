import os
import json
import logging
import re
import requests
from dotenv import load_dotenv
import difflib
from bs4 import BeautifulSoup
import fitz
from azure.core.credentials import AzureKeyCredential
from azure.storage.blob import BlobServiceClient, ContainerClient, ContentSettings  # Import C
from docx import Document
import fitz  # PyMuPDF
from io import BytesIO
from docx import Document
from docx.shared import Pt
from datetime import datetime
import requests
from retrying import retry
load_dotenv()

DB_NAME = os.environ['DB_NAME']
DB_HOST = os.environ['DB_HOST']
DB_PORT = os.environ['DB_PORT']
DB_USER = os.environ['DB_USER']
DB_PASSWORD = os.environ['DB_PASSWORD']
DB_SCHEMA = 'dcompare'
RESUME_ASSISTANCE_ID = os.environ['RESUME_ASSISTANCE_ID']
LEASE_ASSISTANCE_ID = os.environ['LEASE_ASSISTANCE_ID']
EMAIL_ASSISTANCE_ID = os.environ['EMAIL_ASSISTANCE_ID']
OPENAI_API_KEY = os.environ['OPENAI_API_KEY']
EMAIL_ID = os.environ['EMAIL_ID']
PASSWORD = os.environ['PASSWORD']
FILE_TEMPLATE_URL = 'https://crcp.impressicocrm.com/api/api/ResumeParse'
# FILE_TEMPLATE_URL = 'http://localhost:5243/api/ResumeParse'

'''
is_simulating = int(os.environ['IS_SIMULATING'])

abs_path = os.path.abspath(os.path.dirname(__file__))
cred_path = os.path.join(abs_path, 'gcp-credentials.json')
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r'{}'.format(cred_path)
'''
logging.basicConfig(
    format='%(asctime)s %(levelname)-5s %(module)-10s %(message)s',
    level=logging.INFO,
    datefmt='%Y-%m-%d %H:%M:%S')
logger = logging.getLogger(__name__)


def unique_list_keeping_order(seq):
    seen = set()
    seen_add = seen.add
    return [x for x in seq if not (x in seen or seen_add(x))]


def make_dict(fields, data):
    '''
    fields are present in data list in the same order
    '''
    return {fields[i]: data[i] for i in range(len(fields))}


def fields_as(field):
    fs = field.split(' as ')
    return fs[0] if len(fs) == 1 else fs[-1]


def wrap_long_text(s, n_width):
    return [s]


# def blob_client_get(blob_name):
#     from azure.core.credentials import AzureKeyCredential
#     from azure.storage.blob import BlobServiceClient, ContainerClient, ContentSettings  # Import ContentSettings
#     connection_string = os.environ['connection_string']
#     container_name = os.environ['container_name']

#     blob_service_client = BlobServiceClient.from_connection_string(connection_string)
#     container_client = blob_service_client.get_container_client(container_name)

#     # Create a ContainerClient and upload the file
#     blob_client = container_client.get_blob_client(blob_name)
#     return blob_client

def upload_file_to_cloud(file_data, filename, content_type):
    blob_client = blob_client_get(filename)
    # Set cache control for the blob
    content_settings = ContentSettings(content_type=content_type['content_type'], cache_control='no-store')  # Add cache_control

    blob_client.upload_blob(file_data, content_settings=content_settings, overwrite=True)

    # Return the path to the uploaded file
    url = 'https://{blob_service_client.account_name}.blob.core.windows.net/{container_name}/{blob_name}'
    return {'url': url}

def download_file_from_cloud(blob_name, folder):
    blob_client = blob_client_get(blob_name)
    os.makedirs(folder, exist_ok=True)
    fpath = os.path.join(folder, blob_name)
    with open(fpath, 'wb') as local_file:
        blob_data = blob_client.download_blob()
        local_file.write(blob_data.readall())
        return fpath

def extract_table_content(html_diff):
    # Parse the HTML diff
    soup = BeautifulSoup(html_diff, 'html.parser')

    # Find the table element
    table = soup.find('table')

    # If a table is found, return its HTML content, else return None
    if table:
        return str(table)
    else:
        return None


def wrap_in_html_table(table_content):
    if table_content:
        # Wrap the table content in a minimal HTML structure
        html = f"<html><body>{table_content}</body></html>"
        return html
    else:
        return None


def table_diff_in_html(file1_lines, file2_lines, desc1, desc2):
    # Perform the line-level comparison
    file2_lines = [item for item in file2_lines if item != '\n']
    file1_lines = [item for item in file1_lines if item != '\n']

    # Perform the line-level comparison
    differ = difflib.HtmlDiff(wrapcolumn=90, linejunk=None, charjunk=None)
    html_diff = differ.make_file(
        file1_lines, file2_lines, fromdesc=desc1, todesc=desc2, context=True, numlines=50)

    # Extract the table content from the HTML diff
    table_content = extract_table_content(html_diff)

    # Wrap the table content in an HTML structure
    html_table = wrap_in_html_table(table_content)

    return html_table


def add_header(header_txt, paragraph):
    run = paragraph.add_run(f"{header_txt}: \n")
    run.bold = True
    run.font.name = "Calibri"


def add_sub_header(sub_header_txt, paragraph):
    run = paragraph.add_run(f"   {sub_header_txt} ")
    run.bold = True
    run.font.name = "Calibri"


def remove_paragraph_element(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def update_resume_with_template(file_name, resume_txt, template_name):
    # Open the template document
    document = Document(f'{template_name}.docx')
    if template_name == 'blank_template':
        document.add_paragraph(f"{resume_txt}")

        # Save the updated document with a new name to avoid overwriting the template
        document.save(f'{file_name}')
        return

    resume_dict = json.loads(resume_txt)
    resume_info = Resume_Info(resume_dict.get('candidate_name', ''),
                              resume_dict.get('phone_number', ''),
                              resume_dict.get('email', ''),
                              resume_dict.get('career_objective', ''),
                              resume_dict.get('most_recent_job_title', ''),
                              resume_dict.get('academic_qualifications', ''),
                              resume_dict.get(
                                  'non_academic_qualifications', ''),
                              resume_dict.get('computer_skills', ''),
                              resume_dict.get('other_skills', ''),
                              resume_dict.get('employment_history', ''),
                              resume_dict.get('project_description_and_responsibilities', ''))

    # Placeholders to be replaced
    replacements = {
        'candidate_name': resume_info.candidate_name,
        'candidate_phone': resume_info.candidate_phone,
        'candidate_email': resume_info.candidate_email,
        'career_objective': resume_info.career_objective,
        'most_recent_job_title': resume_info.most_recent_job_title,
        '[academic_qualifications]': resume_info.academic_qualifications,
        '[non_academic_qualifications]': resume_info.non_academic_qualifications,
        '[technical_skills]': resume_info.technical_skills,
        '[other_skills]': resume_info.other_skills,
        '[employment_history]': resume_info.employment_history,
        '[description_of_project]': resume_info.project_description
    }

    academic_qualifications = 0
    non_academic_qualifications_cnt = 0
    technical_skills_cnt = 0
    other_skills_cnt = 0

    # Replace placeholders in the document's paragraphs
    for key in replacements:
        for paragraph in document.paragraphs:
            if 'academic_qualifications' in paragraph.text and 'non_academic_qualifications' not in paragraph.text:
                if academic_qualifications == 1:
                    remove_paragraph_element(paragraph)
                # Clear the placeholder line
                paragraph.clear()
                if len(resume_info.academic_qualifications) > 0:
                    add_header("Academic Qualifications", paragraph)
                    for qualification in resume_info.academic_qualifications:
                        paragraph.add_run(
                            f"   • {qualification.get('degree', '')} from {qualification.get('university', '')}, {qualification.get('location', '')}\n")
                        paragraph.add_run(
                            f"     ({qualification.get('duration', '')})\n")
                    academic_qualifications = academic_qualifications + 1
                else:
                    remove_paragraph_element(paragraph)

            elif 'non_academic_qualifications' in paragraph.text:
                if non_academic_qualifications_cnt == 1:
                    remove_paragraph_element(paragraph)
                # Clear the placeholder line
                paragraph.clear()
                if len(resume_info.non_academic_qualifications) > 0:
                    add_header(
                        "Awards and Non-academic Qualifications", paragraph)
                    non_academic_qualifications = ', '.join(
                        resume_info.non_academic_qualifications)
                    paragraph.add_run(f"     {non_academic_qualifications}\n")
                    non_academic_qualifications_cnt = non_academic_qualifications_cnt + 1
                else:
                    remove_paragraph_element(paragraph)
            elif 'technical_skills' in paragraph.text:
                if technical_skills_cnt == 1:
                    remove_paragraph_element(paragraph)
                # Clear the placeholder line
                paragraph.clear()
                if len(resume_info.technical_skills) > 0:
                    add_header("Technical Skills", paragraph)
                    technical_skills = ', '.join(resume_info.technical_skills)
                    paragraph.add_run(f"     {technical_skills}\n")
                    technical_skills_cnt = technical_skills_cnt + 1
                else:
                    remove_paragraph_element(paragraph)
            elif 'other_skills' in paragraph.text:
                if other_skills_cnt == 1:
                    continue
                # Clear the placeholder line
                paragraph.clear()
                if len(resume_info.other_skills) > 0:
                    add_header("Other Skills", paragraph)
                    other_skills = ', '.join(resume_info.other_skills)
                    paragraph.add_run(f"     {other_skills}\n")
                    other_skills_cnt = other_skills_cnt + 1
                else:
                    remove_paragraph_element(paragraph)
            elif 'employment_history' in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, "")
                # Clear the placeholder line
                paragraph.clear()
                if len(resume_info.employment_history) > 0:
                    add_header("Employment History", paragraph)
                    for employment in resume_info.employment_history:
                        add_sub_header(
                            f"{employment.get('designation', '')}", paragraph)
                        paragraph.add_run(
                            f" {employment.get('duration', '')}\n")
                        paragraph.add_run(
                            f"   {employment.get('organization', '')}\n")
                        paragraph.add_run("\n")
                else:
                    remove_paragraph_element(paragraph)
            elif 'project_description' in paragraph.text:
                # Clear the placeholder line
                paragraph.clear()
                if len(resume_info.project_description) > 0:
                    add_header(
                        "Project Description And Responsibilities", paragraph)
                    for project in resume_info.project_description:
                        add_sub_header("Project Name:", paragraph)
                        paragraph.add_run(
                            f"{project.get('project_name', '')}\n")

                        add_sub_header("Description:", paragraph)
                        paragraph.add_run(
                            f"{project.get('project_description', '')}\n")

                        if len(project.get('responsibilities', '')) > 0:
                            responsibilities_date_run = paragraph.add_run(
                                f"   Responsibilities: \n")
                            responsibilities_date_run.bold = True
                            responsibilities_date_run.font.name = "Calibri"
                            paragraph.add_run(
                                f"      {project.get('responsibilities', '')}\n")

                        paragraph.add_run("\n")
                else:
                    remove_paragraph_element(paragraph)
            else:
                for run in paragraph.runs:
                    replace_placeholder_with_formatting(
                        run, key, replacements.get(key, ''))
    for paragraph in document.paragraphs:
        if "Project Name: project_name" in paragraph.text\
                or "Responsibilities: project_responsibilities" in paragraph.text \
                or "job_title duration" in paragraph.text \
                or "company" == paragraph.text \
                or "other_skills_value" == paragraph.text \
                or "description_of_project" in paragraph.text:
            remove_paragraph_element(paragraph)
    # Save the updated document with a new name to avoid overwriting the template
    document.save(f'{file_name}')


@retry(wait_fixed=2000, stop_max_attempt_number=3)
def call_post_api(url, data):
    response = requests.post(url, json=data)
    # Raise an exception for 4xx or 5xx status codes
    response.raise_for_status()
    return response.content, response.status_code


def create_folder_not_exist(folder_name):
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
        print(f"Folder '{folder_name}' created successfully.")
    else:
        print(f"Folder '{folder_name}' already exists.")


def update_lease_agreement_with_template(json_data, template_name, new_name):
    data = json.loads(json_data)
    data["documentName"] = template_name
    data["documentType"] = 2

    content, status_code = call_post_api(url=FILE_TEMPLATE_URL, data=data)
    if status_code == 200:
        with open(new_name, 'wb') as f:
            f.write(content)
    return status_code


def replace_placeholder_with_formatting(run, placeholder, replacement):
    if placeholder in run.text:
        if isinstance(replacement, list):
            print('Replacement is a List')
        else:
            run.text = run.text.replace(placeholder, replacement)


def delete_file(file_name):
    try:
        os.remove(file_name)
    except:
        logger.info('File does not exist')


class Resume_Info:
    def __init__(self,
                 candidate_name,
                 candidate_phone,
                 candidate_email,
                 career_objective,
                 most_recent_job_title,
                 academic_qualifications,
                 non_academic_qualifications,
                 technical_skills,
                 other_skills,
                 employment_history,
                 project_description):
        self.candidate_name = candidate_name
        self.candidate_phone = candidate_phone
        self.candidate_email = candidate_email
        self.career_objective = career_objective
        self.most_recent_job_title = most_recent_job_title
        self.academic_qualifications = academic_qualifications
        self.non_academic_qualifications = non_academic_qualifications
        self.technical_skills = technical_skills
        self.other_skills = other_skills
        self.employment_history = employment_history
        self.project_description = project_description


def remove_prefix(text, prefix):
    if text.startswith(prefix):
        return text[len(prefix):]
    else:
        return text


def remove_suffix(text, suffix):
    if text.endswith(suffix):
        return text[:-len(suffix)]
    else:
        return text


def fetch_emails_and_categorize():
    import imaplib
    import email
    from email.header import decode_header
    from mydb import run_select, run_idu
    from aihelper import get_resume_categories

    # Login to IMAP server
    # Change this to your IMAP server
    imap = imaplib.IMAP4_SSL('pop.gmail.com')
    imap.login(EMAIL_ID, PASSWORD)
    # Select a mailbox (INBOX, for instance)
    imap.select('inbox')
    # Search emails
    # You might want to narrow down this search
    status, messages = imap.search(None, 'ALL')
    resume_categories = []
    for num in messages[0].split():
        # Fetch the email by ID
        typ, msg_data = imap.fetch(num, '(RFC822)')
        msg = {}
        body = ''
        attachment_content = ''
        subject = ''
        from_email = None
        filenames = []
        timedate = None
        categories_dict = {}

        for response_part in msg_data:
            if isinstance(response_part, tuple):
                msg = email.message_from_bytes(response_part[1])
                date = msg.get("Date")
                epoch_time, timedate = convert_str_epoch(date)
                filename_with_epoch = epoch_time + ".docx"
                rows = run_select(f"select attachment_name from {DB_SCHEMA}.email_data where cloud_file_name=%s",
                                  (filename_with_epoch,))
                from_email = msg.get("From")
                result = rows[0][0] if rows else None
                if result is None:
                    # Decode email subject
                    subject = decode_header(msg["Subject"])[0][0]
                    if isinstance(subject, bytes):
                        subject = subject.decode()
                    # Check if the email message is multipart
                    if msg.is_multipart():
                        for part in msg.walk():
                            # Get the content type of the email part
                            content_type = part.get_content_type()
                            content_disposition = str(
                                part.get("Content-Disposition"))
                            # Ignore attachments/html
                            if content_type == "text/plain" and "attachment" not in content_disposition:
                                body = part.get_payload(decode=True).decode()
                            if "attachment" in content_disposition:
                                # this part is the attachment
                                filename = part.get_filename()
                                filenames.append(filename)
                                if filename.endswith('.docx') or filename.endswith('.pdf'):
                                    attachment_bytes = part.get_payload(
                                        decode=True)
                                if filename.endswith('.docx'):
                                    # Read DOCX content from bytes
                                    doc = Document(BytesIO(attachment_bytes))
                                    doc_text = "\n".join(
                                        [para.text for para in doc.paragraphs])
                                    attachment_content = doc_text + "\n " + attachment_content
                                elif filename.endswith('.pdf'):
                                    # Read PDF content from bytes
                                    pdf_file = fitz.open(
                                        stream=attachment_bytes, filetype="pdf")
                                    pdf_text = ""
                                    for page_num in range(len(pdf_file)):
                                        page = pdf_file.load_page(page_num)
                                        pdf_text += page.get_text()
                                    attachment_content = pdf_text + "\n " + attachment_content
                    else:
                        # If not multipart, just extract the body
                        body = msg.get_payload(decode=True).decode()
                    doc = Document()
                    prepare_doc("Email Subject:", subject, doc)
                    prepare_doc("Email Body:", body, doc)
                    prepare_doc("Attachment Name:", filenames[0], doc)
                    prepare_doc("Attachment Content:", attachment_content, doc)
                    content_type = (
                        {"content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})
                    document_content = ""
                    # Iterate through each paragraph in the document and concatenate their text
                    for paragraph in doc.paragraphs:
                        document_content += paragraph.text + "\n"
                    res = upload_file_to_cloud(
                        document_content, filename_with_epoch, content_type)
                    cloud_file_url = res.get("url", "")
                    download_file_from_cloud(
                        filename_with_epoch, "email-parser")
                    email_txt = f"Subject: {subject}" + "\n" + f"Body: {body}" + \
                        "\n" + f"Attachment content: {attachment_content}"
                    categories_dict = get_resume_categories(email_txt)
                    run_idu("""insert into dcompare.email_data(subject, body, attachment_name,attachment_content,
                    cloud_file_url,cloud_file_name,location,experience_level,role,received_date, sender_email)
                    values (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""", (
                        subject, body, filenames[0], attachment_content, cloud_file_url, filename_with_epoch,
                        categories_dict['location'],  categories_dict['experience_level'], categories_dict['role'], timedate, from_email))
                    doc_upload_to_assistance(
                        "email-parser/" + filename_with_epoch)
                else:
                    data = get_email_data(filename_with_epoch)
                    body = data['body']
                    attachment_content = data['attachment_content']
                    filenames.append(data['attachment_name'])
                    subject = data['subject']
                    categories_dict['location'] = data['location']
                    categories_dict['experience_level'] = data['experience_level']
                    categories_dict['role'] = data['role']
        categories_dict['from_email'] = from_email
        categories_dict['date'] = timedate
        categories_dict['subject'] = subject
        categories_dict['body'] = body
        categories_dict['attachments'] = filenames
        categories_dict['attachment_content'] = attachment_content
        resume_categories.append(categories_dict)
    return resume_categories


def get_email_data(filename_with_epoch):
    from mydb import run_select
    fields = [
        "subject",
        "body",
        "attachment_name",
        "attachment_content",
        "location",
        "experience_level",
        "role"
    ]
    s_fields = ",".join(fields)
    query = f"""select {s_fields} from dcompare.email_data where cloud_file_name=%s"""
    result = run_select(query, (filename_with_epoch,))
    fields2 = [fields_as(field) for field in fields]
    d = make_dict(fields2, result[0]) if result else {}
    return d


def prepare_doc(header, body, doc):
    doc.add_paragraph().add_run(header).bold = True
    doc.add_paragraph().add_run(body)


def doc_upload_to_assistance(file_path):
    url = "https://api.openai.com/v1/files"

    # Prepare headers with Authorization token
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}"
    }

    # Prepare payload
    payload = {
        "purpose": "assistants"
    }

    # Prepare files
    files = {
        "file": open(file_path, "rb")
    }

    # Send the POST request
    response = requests.post(url, headers=headers, data=payload, files=files)

    # Check the response
    if response.status_code == 200:
        # Parse response JSON
        response_json = response.json()
        # Extract file_id
        file_id = response_json.get("id")
        if file_id:
            create_assistance_file(file_id)
        else:
            print("Failed to read file_id from response.")
        print("File uploaded successfully.")
    else:
        print("Failed to upload file. Error:", response.text)


def create_assistance_file(file_id):
    # API endpoint
    url = "https://api.openai.com/v1/assistants/{}/files".format(
        EMAIL_ASSISTANCE_ID)

    # Request headers
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json",
        "OpenAI-Beta": "assistants=v1"
    }

    # Request body
    data = {
        "file_id": file_id
    }

    # Send the POST request
    response = requests.post(url, headers=headers, json=data)

    # Check the response
    if response.status_code == 200:
        print("File linked to assistant successfully.")
    else:
        print("Failed to link file to assistant. Error:", response.text)


def convert_str_epoch(date_str):
    date_format = "%a, %d %b %Y %H:%M:%S %z"

    # Convert string to datetime object
    dt_object = datetime.strptime(date_str, date_format)

    # Convert datetime object to epoch time in milliseconds
    epoch_time_milliseconds = int(dt_object.timestamp() * 1000)

    return str(epoch_time_milliseconds), str(dt_object)


def response(status, s, body, code):
    return {"status": status, "msg": s, "body": body}, code, {'Content-Type': 'application/json'}


def download_from_cloud(filename, directory_name):
    if not os.path.exists(directory_name+'/'+filename):
        download_file_from_cloud(filename, directory_name)
