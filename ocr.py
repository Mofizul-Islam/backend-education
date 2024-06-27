from docx import Document
import re
import subprocess
#from pdf2image import convert_from_path
import tempfile
import fitz  # PyMuPDF


def separate_cells(cell_texts):
    return "\n-----------------------------------------------------------------------\n".join(
        cell_texts
    )


"""Extract text from the DOCX file"""


def extract_text_from_docx(filename):
    doc = Document(filename)

    """Extract text from tables"""
    english_texts = []
    croatian_texts = []
    for table in doc.tables:
        for row in table.rows:
            """Ensure there are at least two cells in the row"""
            if len(row.cells) >= 2:
                english_texts.append(row.cells[0].text)
                croatian_texts.append(row.cells[1].text)

    separated_english_texts = separate_cells(english_texts)
    separated_croatian_texts = separate_cells(croatian_texts)

    return (
        separated_english_texts,
        separated_croatian_texts,
    )


# def Extract_pdf(file_name):
#     dd_data = ""
#     images = convert_from_path(file_name)
#     for _, image in enumerate(images):
#         with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as temp_file:
#             image_path = temp_file.name
#             image.save(image_path, "JPEG")
#             # data = your_funtion (image)
#             pass


def extract_text_from_pdf(pdf_path):
    # Open the PDF file
    document = fitz.open(pdf_path)
    text = ""

    # Iterate over each page
    for page_num in range(len(document)):
        page = document.load_page(page_num)  # Load page
        text += page.get_text()  # Extract text from page

    return text

def get_docx_text(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

if __name__ == "__main__":
    filename = '/Users/ruchirgarg/Downloads/anil-resume.docx'
    print(get_docx_text(filename))

